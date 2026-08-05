"""
services/export_service.py
===========================
Generates exportable representations of translation records in several
formats. Each method returns raw bytes ready to hand to
`st.download_button`, so this service has no Streamlit dependency itself.
"""

from __future__ import annotations

import csv
import io
import json

from utils.exceptions import ExportError
from utils.logger import get_logger

logger = get_logger(__name__)


class ExportService:
    """Builds export files from a list of history-record dicts (as
    returned by HistoryService)."""

    def to_txt(self, records: list[dict]) -> bytes:
        lines = []
        for r in records:
            lines.append(f"[{r['created_at']}] {r['source_language']} -> {r['target_language']}")
            lines.append(f"Original: {r['source_text']}")
            lines.append(f"Translated: {r['translated_text']}")
            lines.append("-" * 40)
        return "\n".join(lines).encode("utf-8")

    def to_csv(self, records: list[dict]) -> bytes:
        if not records:
            return b""
        buffer = io.StringIO()
        fieldnames = [
            "id", "created_at", "source_language", "target_language",
            "source_text", "translated_text", "model_used",
            "word_count", "character_count", "is_favorite",
        ]
        writer = csv.DictWriter(buffer, fieldnames=fieldnames, extrasaction="ignore")
        writer.writeheader()
        for r in records:
            row = dict(r)
            row["created_at"] = str(row["created_at"])
            writer.writerow(row)
        return buffer.getvalue().encode("utf-8")

    def to_json(self, records: list[dict]) -> bytes:
        serializable = []
        for r in records:
            row = dict(r)
            row["created_at"] = str(row["created_at"])
            serializable.append(row)
        return json.dumps(serializable, indent=2, ensure_ascii=False).encode("utf-8")

    def to_docx(self, records: list[dict]) -> bytes:
        try:
            from docx import Document
        except ImportError as exc:
            raise ExportError(
                f"python-docx not installed: {exc}",
                user_message="DOCX export is not available right now.",
            ) from exc

        try:
            doc = Document()
            doc.add_heading("ALT Translation Export", level=1)
            for r in records:
                doc.add_heading(f"{r['source_language']} → {r['target_language']}", level=3)
                doc.add_paragraph(f"Date: {r['created_at']}")
                doc.add_paragraph(f"Original: {r['source_text']}")
                doc.add_paragraph(f"Translated: {r['translated_text']}")
                doc.add_paragraph("")

            buffer = io.BytesIO()
            doc.save(buffer)
            return buffer.getvalue()
        except Exception as exc:  # noqa: BLE001
            raise ExportError(
                f"DOCX generation failed: {exc}",
                user_message="Could not generate the Word document.",
            ) from exc

    def to_pdf(self, records: list[dict]) -> bytes:
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet
        except ImportError as exc:
            raise ExportError(
                f"reportlab not installed: {exc}",
                user_message="PDF export is not available right now.",
            ) from exc

        try:
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter)
            styles = getSampleStyleSheet()
            story = [Paragraph("ALT Translation Export", styles["Title"]), Spacer(1, 12)]

            for r in records:
                story.append(Paragraph(f"{r['source_language']} → {r['target_language']} ({r['created_at']})", styles["Heading3"]))
                story.append(Paragraph(f"Original: {_escape(r['source_text'])}", styles["Normal"]))
                story.append(Paragraph(f"Translated: {_escape(r['translated_text'])}", styles["Normal"]))
                story.append(Spacer(1, 12))

            doc.build(story)
            return buffer.getvalue()
        except Exception as exc:  # noqa: BLE001
            raise ExportError(
                f"PDF generation failed: {exc}",
                user_message="Could not generate the PDF.",
            ) from exc


def _escape(text: str) -> str:
    """Escape characters that would otherwise be interpreted as markup
    by reportlab's Paragraph (which parses a small HTML-like subset)."""
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )
